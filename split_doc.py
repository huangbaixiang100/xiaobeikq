#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档分割工具 - 从源代码文档中提取指定页面范围
功能：
1. 提取前30页和后30页组成新文档
2. 提取前25页和后25页组成新文档
"""

from docx import Document
from docx.shared import Inches, Mm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement, parse_xml
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.shared import qn
from copy import deepcopy
import os
import shutil

def copy_section_properties(source_section, target_section):
    """复制文档节属性（页面设置等）"""
    target_section.page_height = source_section.page_height
    target_section.page_width = source_section.page_width
    target_section.left_margin = source_section.left_margin
    target_section.right_margin = source_section.right_margin
    target_section.top_margin = source_section.top_margin
    target_section.bottom_margin = source_section.bottom_margin
    target_section.header_distance = source_section.header_distance
    target_section.footer_distance = source_section.footer_distance

def copy_header(source_doc, target_doc):
    """复制页眉设置"""
    source_section = source_doc.sections[0]
    target_section = target_doc.sections[0]
    
    # 复制页眉
    source_header = source_section.header
    target_header = target_section.header
    
    # 清空目标页眉的现有内容
    for p in list(target_header.paragraphs):
        p._element.getparent().remove(p._element)
    
    # 添加新的页眉段落
    new_p = target_header.add_paragraph()
    new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加软件名称和版本号（从第一个run获取）
    first_p = source_header.paragraphs[0]
    text_parts = first_p.text.split("  ")  # 假设用两个空格分隔
    if len(text_parts) >= 2:
        software_name = text_parts[0]
        version = text_parts[1]
        
        # 添加软件名称
        run = new_p.add_run(f"{software_name}  ")
        run.font.size = Pt(10)
        run.font.name = 'Courier New'
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        
        # 添加版本号
        run = new_p.add_run(f"{version}  ")
        run.font.size = Pt(10)
        run.font.name = 'Courier New'
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    
    # 添加页码
    run = new_p.add_run("第")
    run.font.size = Pt(10)
    run.font.name = 'Courier New'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    
    # 插入页码域代码
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    new_p.runs[-1]._r.append(fldChar_begin)
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = " PAGE "
    new_p.runs[-1]._r.append(instrText)
    
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    new_p.runs[-1]._r.append(fldChar_end)
    
    # 添加"页"字
    run = new_p.add_run("页")
    run.font.size = Pt(10)
    run.font.name = 'Courier New'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

def extract_lines(source_path, output_path, front_lines, back_lines):
    """
    从源文档中提取指定数量的前后行数
    
    Args:
        source_path: 源文档路径
        output_path: 输出文档路径
        front_lines: 要提取的前几行
        back_lines: 要提取的后几行
    """
    # 加载源文档
    source_doc = Document(source_path)
    target_doc = Document()
    
    # 复制文档属性
    copy_section_properties(source_doc.sections[0], target_doc.sections[0])
    copy_header(source_doc, target_doc)
    
    # 收集所有行
    all_lines = []
    for table in source_doc.tables:
        for row in table.rows:
            cell = row.cells[0]
            if cell.paragraphs:
                text = cell.paragraphs[0].text
                all_lines.append(text)
    
    total_lines = len(all_lines)
    print(f"源文档共有 {total_lines} 行")
    
    # 确保不超出实际行数
    front_lines = min(front_lines, total_lines)
    back_lines = min(back_lines, total_lines - front_lines)
    
    # 准备要写入的行
    lines_to_write = []
    # 添加前面的行
    lines_to_write.extend(all_lines[:front_lines])
    # 添加后面的行
    if back_lines > 0:
        lines_to_write.extend(all_lines[-back_lines:])
    
    # 按每页50行分页写入
    lines_per_page = 50
    total_pages = (len(lines_to_write) + lines_per_page - 1) // lines_per_page
    
    for page in range(total_pages):
        start_idx = page * lines_per_page
        end_idx = min(start_idx + lines_per_page, len(lines_to_write))
        page_lines = lines_to_write[start_idx:end_idx]
        
        # 创建表格
        table = target_doc.add_table(rows=len(page_lines), cols=1)
        table.style = 'Table Grid'
        table.autofit = False
        
        # 写入行
        for i, line in enumerate(page_lines):
            cell = table.rows[i].cells[0]
            p = cell.paragraphs[0]
            p.text = line
            # 设置单元格格式
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            # 设置段落格式
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.line_spacing = Pt(12)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            # 设置字体
            run = p.runs[0] if p.runs else p.add_run()
            run.font.name = 'Courier New'
            run.font.size = Pt(8)
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
            
        # 添加分页符（除了最后一页）
        if page < total_pages - 1:
            target_doc.add_page_break()
    
    # 保存新文档
    target_doc.save(output_path)
    print(f"已生成文档：{output_path}")
    print(f"包含前 {front_lines} 行和后 {back_lines} 行，共 {total_pages} 页")

def main():
    # 源文档路径
    source_path = r"C:\Users\柏\Desktop\源代码文档.docx"
    
    # 确保源文档存在
    if not os.path.exists(source_path):
        print(f"错误：找不到源文档 {source_path}")
        return
    
    # 输出目录
    output_dir = os.path.dirname(source_path)
    
    # 生成两个版本的文档
    # 1. 前1500行和后1500行（生成60页，每页50行）
    output_path_30 = os.path.join(output_dir, "前30后30.docx")
    extract_lines(source_path, output_path_30, 1500, 1500)
    
    # 2. 前1250行和后1250行（生成50页，每页50行）
    output_path_25 = os.path.join(output_dir, "前25后25.docx")
    extract_lines(source_path, output_path_25, 1250, 1250)

if __name__ == "__main__":
    main()
