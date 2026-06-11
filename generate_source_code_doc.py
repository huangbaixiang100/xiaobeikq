
"""
源代码整理工具- 用于软著申请
页眉包含：软件名  版本号  第{PAGE}页
"""

import os
import sys
import chardet  
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.shared import OxmlElement, qn
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement as Ox, parse_xml
from typing import List, Iterable, Optional
import re



def detect_encoding(path: str, default: str = "utf-8"):

    try:
        with open(path, "rb") as f:
            raw = f.read(4096)
        if raw:
            enc = chardet.detect(raw).get("encoding") if 'chardet' in sys.modules else None
            if enc:
                return enc
    except Exception:
        pass
   
    return default

def read_text_safely(path: str) -> str:
    """读取文本"""
    encodings = []
    encodings.append(detect_encoding(path))
    encodings += ['utf-8', 'utf-8-sig', 'gbk', 'gb2312', 'latin-1']
    tried = set()
    for enc in encodings:
        if not enc or enc.lower() in tried:
            continue
        tried.add(enc.lower())
        try:
            with open(path, 'r', encoding=enc, errors='strict') as f:
                return f.read()
        except Exception:
            continue
 
    with open(path, 'r', encoding='latin-1', errors='ignore') as f:
        return f.read()

def remove_empty_lines(text: str) -> str:
    """移除空行/纯空白行"""
    out = []
    for line in text.splitlines():
        if line.strip(): 
            out.append(line.rstrip())
    return "\n".join(out)

def iter_files(root_or_files: Iterable[str],
               allow_exts: Optional[Iterable[str]] = None) -> List[str]:

    if allow_exts is None:
        allow_exts = {
            ".py", ".js", ".ts", ".tsx", ".jsx", ".html", ".css", ".json",
            ".yml", ".yaml", ".toml", ".md", ".txt", ".java", ".c", ".cpp",
            ".h", ".hpp", ".cs", ".go", ".rs", ".php", ".rb", ".sh", ".bat"
        }
    files = []
    def add_file(p: str):
        if os.path.isfile(p):
            if os.path.splitext(p)[1].lower() in allow_exts:
                files.append(os.path.abspath(p))

    for item in root_or_files:
        p = os.path.abspath(item)
        if os.path.isdir(p):
            for dirpath, _, filenames in os.walk(p):
                for fn in sorted(filenames):
                    add_file(os.path.join(dirpath, fn))
        else:
            add_file(p)
    return sorted(files)

# ----------------------------
# Word 样式和页眉
# ----------------------------

def add_header_with_info(doc: Document, software_name: str, version: str):
    """页眉：软件名  版本号  第{PAGE}页"""
    section = doc.sections[0]
    header = section.header

  
    for p in header.paragraphs:
        for r in p.runs:
            r.clear()

    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_text_run(text):
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.name = 'Courier New'  
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        return run

    add_text_run(f"{software_name}  ")
    add_text_run(f"{version}  ")
    add_text_run("第")


    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    p.runs[-1]._r.append(fldChar_begin)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = " PAGE "
    p.runs[-1]._r.append(instrText)

    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    p.runs[-1]._r.append(fldChar_end)

    add_text_run("页")

def ensure_code_paragraph_style(paragraph):
    """将段落设置为等宽 + 固定行距 + 零间距"""
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(12)          # 固定 12pt 行高
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.keep_together = True           # 强制整段同页
    pf.keep_with_next = True          # 与下一段同页
    pf.widow_control = True           # 避免孤行
    
    # 设置段落属性
    p = paragraph._element
    p.set(qn('w:spacing'), '0')       # 段落间距为0
    p.set(qn('w:ind'), '0')           # 缩进为0
    
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    return paragraph


def add_code_page_table(doc: Document, page_lines: List[str]):
    """
    将一页的代码（最多 50 行）渲染为一个单列表格（每行一格），
    严格控制行距与页内排布，避免 PDF 导出出现视觉"间断"。
    """
    rows = len(page_lines)
    table = doc.add_table(rows=rows, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False  
    table.allow_autofit = False  
    tbl = table._tbl
    tblPr = tbl.tblPr

    tblCellMar = Ox('w:tblCellMar')
    for side in ['top', 'left', 'bottom', 'right']:
        mar = Ox(f'w:{side}')
        mar.set(qn('w:w'), '0')
        mar.set(qn('w:type'), 'dxa')
        tblCellMar.append(mar)
    tblPr.append(tblCellMar)
    
    # 设置表格不允许跨页分割
    tblPr.append(parse_xml(f'<w:tblLook xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:firstRow="0" w:lastRow="0" w:firstColumn="0" w:lastColumn="0" w:noHBand="0" w:noVBand="0"/>'))
    
    # 固定行高（12pt = 240 twips）
    trHeight = Ox('w:trHeight')
    trHeight.set(qn('w:val'), '240')
    trHeight.set(qn('w:hRule'), 'exact')
    
    # 设置表格不允许跨页分割
    tblPr.append(parse_xml('<w:cantSplit xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="1"/>'))
    
    # 设置表格无边框
    tblBorders = parse_xml('''
        <w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:top w:val="none"/>
            <w:left w:val="none"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
            <w:insideH w:val="none"/>
            <w:insideV w:val="none"/>
        </w:tblBorders>
    ''')
    tblPr.append(tblBorders)
    
    # 设置表格行间距为0
    tblPr.append(parse_xml('''
        <w:tblCellSpacing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:w="0" w:type="dxa"/>
    '''))

    for i, line in enumerate(page_lines):
        row = table.rows[i]
        # 应用固定行高
        row._tr.append(trHeight)
        
        cell = row.cells[0]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 改为居中对齐
        # 清除默认段落
        for rp in list(cell.paragraphs):
            p_el = rp._element
            p_el.getparent().remove(p_el)

        p = cell.add_paragraph(line)
        ensure_code_paragraph_style(p)
        # 确保段落不会自动分页
        p._p.set('keepNext', '1')

    # 添加一个空段落，高度为0
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    
    return table


def sanitize_code(line: str) -> str:
    """清理代码中的敏感信息"""
    # 替换IP地址
    line = re.sub(r'\d+\.\d+\.\d+\.\d+(?::\d+)?', '[IP]', line)
    # 替换URL/域名
    line = re.sub(r'https?://[^\s<>"]+|www\.[^\s<>"]+', '[URL]', line)
    # 替换文件路径
    line = re.sub(r'[A-Za-z]:\\[^\s<>"|?*]+', '[PATH]', line)
    line = re.sub(r'/[^\s<>"|?*]+/', '[PATH]/', line)
    # 替换API密钥格式的字符串
    line = re.sub(r'[A-Za-z0-9_-]{30,}', '[KEY]', line)
    # 替换端口号
    line = re.sub(r':\d{2,5}(?=/|$)', ':[PORT]', line)
    # 替换邮箱地址
    line = re.sub(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', '[EMAIL]', line)
    # 替换数据库连接字符串
    line = re.sub(r'(?i)(?:mongodb|mysql|postgresql|redis)://[^\s<>"]+', '[DB_URL]', line)
    # 替换JWT或其他令牌
    line = re.sub(r'(?i)(?:bearer|token|jwt)\s+[a-zA-Z0-9._-]+', '[TOKEN]', line)
    return line

def create_source_code_document(
    file_paths: Iterable[str],
    output_path: str = "源代码文档.docx",
    lines_per_page: int = 50,
    software_name: str = "小雅医生_乳牙滞留图像智能识别系统",
    version: str = "V1.0",
    max_line_len: int = 100
):
    doc = Document()

    # 页面：A4，1 英寸边距（符合常见软著排版）
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        # 强制 A4 纸张大小
        from docx.shared import Mm
        section.page_height = Mm(297)
        section.page_width = Mm(210)

    add_header_with_info(doc, software_name, version)

    # 收集代码行
    all_lines = []
    line_no = 1

    files = iter_files(file_paths)  # 支持传目录/文件混合
    if not files:
        print("未找到待处理的源文件。")
        return

    print(f"将处理 {len(files)} 个文件")
    for fp in files:
        print(f"读取：{os.path.basename(fp)}")  # 只显示文件名，不显示完整路径
        try:
            content = read_text_safely(fp)
        except Exception as e:
            content = f"# 文件读取错误: {e}"

        content = remove_empty_lines(content)
        for raw in content.splitlines():
            line = sanitize_code(raw)  # 清理敏感信息
            if max_line_len and len(line) > max_line_len:
                line = line[:max_line_len - 3] + "..."
            all_lines.append(f"{line_no:4d}  {line}")
            line_no += 1


    total = len(all_lines)
    lines_per_page = 50
    
    full_pages = total // lines_per_page
    remaining_lines = total % lines_per_page
    
    current_pos = 0
    for page in range(full_pages):
        page_lines = all_lines[current_pos:current_pos + lines_per_page]
        add_code_page_table(doc, page_lines)
        doc.add_page_break()
        current_pos += lines_per_page
    
    if remaining_lines > 0:
        page_lines = all_lines[current_pos:]
        add_code_page_table(doc, page_lines)

    doc.save(output_path)
    print(f"文档已保存：{output_path}")
    print(f"共 {total} 行，约 {full_pages} 页（每页 {lines_per_page} 行）")


def main():
    current_dir = os.path.dirname(os.path.abspath(__file__))
    code_extensions = ['.py', '.js', '.html', '.css', '.json']
    
    targets = [
        r"",
    ]
    
    # 输出文件路径
    output_path = os.path.join(current_dir, "源代码文档.docx")
    software_name = "小雅医生_乳牙滞留图像智能识别系统"
    version = "V1.0"

    print("开始生成源代码文档...")
    create_source_code_document(
        targets,
        output_path=output_path,
        lines_per_page=50,
        software_name=software_name,
        version=version,
        max_line_len=100
    )
    print("完成！")

if __name__ == "__main__":
    main()
